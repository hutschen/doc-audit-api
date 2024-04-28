# Copyright (C) 2024 Helmar Hutschenreuter
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

import docx
import pytest

from docaudit.ml.components import (
    DocxParser,
    DocxParserError,
    iter_sources,
    recursively_merge_dicts,
)


@pytest.mark.parametrize(
    "levels",
    [
        [1, 2, 3],
        [3, 2, 1],
        [1, 1, 1],
        # When a paragraph is not a heading, parse_level should return None
        [1, 2, None],
        [None, 1, 2],
        [3, None, 1],
    ],
)
def test_parse_levels(levels):
    # Create a new document with headings/paragraphs of the specified levels
    doc = docx.Document()
    for level in levels:
        if level is None:
            doc.add_paragraph("Some text")
        else:
            doc.add_heading(f"Heading {level}", level)

    # Parse the document and check if the levels are correct
    for paragraph, level in zip(doc.paragraphs, levels):
        assert DocxParser.parse_level(paragraph) == level


def test_open_docx_carefully_success(monkeypatch):
    # Patch the docx.Document constructor to return a mock document
    mock_document = object()

    def get_mock_document(*args, **kwargs):
        return mock_document

    monkeypatch.setattr(docx, "Document", get_mock_document)

    # Call the open_docx_carefully method and check if it returns the mock document
    assert DocxParser.open_docx_carefully("valid.docx") is mock_document


@pytest.mark.parametrize(
    "paragraphs, expected_results",
    [
        # Case 1: Empty document
        ([], []),
        # Case 2: Document with a single paragraph
        ([("Content 1", None)], [((), "Content 1")]),
        # Case 3: Document with a single heading
        ([("Heading 1", 1)], [(("Heading 1",), "Heading 1")]),
        # Case 4: Document with a single heading and further paragraphs
        (
            [("Heading 1", 1), ("Content 1", None), ("Content 2", None)],
            [(("Heading 1",), "Heading 1\n\nContent 1\n\nContent 2")],
        ),
        # Case 5: Document with multiple nested headings and paragraphs
        (
            [
                ("Heading 1", 1),
                ("Content 1", None),
                ("Heading 2", 2),
                ("Content 2", None),
                ("Content 3", None),
                ("Heading 3", 3),
                ("Content 4", None),
            ],
            [
                (("Heading 1",), "Heading 1\n\nContent 1"),
                (("Heading 1", "Heading 2"), "Heading 2\n\nContent 2\n\nContent 3"),
                (("Heading 1", "Heading 2", "Heading 3"), "Heading 3\n\nContent 4"),
            ],
        ),
        # Case 6: Document with multiple headings at the same level
        (
            [("Heading 1", 1), ("Heading 2", 1), ("Heading 3", 1)],
            [
                (("Heading 1",), "Heading 1"),
                (("Heading 2",), "Heading 2"),
                (("Heading 3",), "Heading 3"),
            ],
        ),
        # Case 7: Document with multiple headings at different levels
        (
            [("Heading 1", 1), ("Heading 2", 3), ("Heading 3", 2)],
            [
                (("Heading 1",), "Heading 1"),
                (("Heading 1", "Heading 2"), "Heading 2"),
                (("Heading 1", "Heading 3"), "Heading 3"),
            ],
        ),
    ],
)
def test_parse(monkeypatch, paragraphs, expected_results):
    # Patch the open_docx_carefully method to return a document with the specified
    # paragraphs
    def mock_open_docx_carefully(*args, **kwargs):
        doc = docx.Document()
        for text, level in paragraphs:
            if level is None:
                doc.add_paragraph(text)
            else:
                doc.add_heading(text, level)
        return doc

    monkeypatch.setattr(DocxParser, "open_docx_carefully", mock_open_docx_carefully)

    # Call the parse method and check if the results are as expected
    assert list(DocxParser.parse("valid.docx")) == expected_results


def test_open_docx_carefully_failure(monkeypatch):
    # Patch the docx.Document constructor to raise an exception
    error_msg = "This is a test error message"

    def mock_document(*args, **kwargs):
        raise Exception(error_msg)

    monkeypatch.setattr(docx, "Document", mock_document)

    # Call the open_docx_carefully method check for DocxParserError
    with pytest.raises(DocxParserError) as exc_info:
        DocxParser.open_docx_carefully("invalid.docx")

    # Assert that the exception contains the error message from the raised exception
    assert error_msg in str(exc_info.value)


@pytest.mark.parametrize(
    "sources, source_ids, expected",
    [
        # Case 1: If source_ids is None
        (["Source1", "Source2"], None, [("Source1", "UUID"), ("Source2", "UUID")]),
        # Case 2: If source_ids is shorter than sources
        (
            ["Source1", "Source2", "Source3"],
            ["ID1"],
            [("Source1", "ID1"), ("Source2", "UUID"), ("Source3", "UUID")],
        ),
        # Case 3: If source_ids is the same length as sources
        (
            ["Source1", "Source2"],
            ["ID1", "ID2"],
            [("Source1", "ID1"), ("Source2", "ID2")],
        ),
        # Case 4: If source_ids is longer than sources
        (["Source1"], ["ID1", "ID2", "ID3"], [("Source1", "ID1")]),
    ],
)
def test_iter_sources(monkeypatch, sources, source_ids, expected):
    # Patch function that returns predictable dummy UUIDs
    def mock_new_source_id():
        return "UUID"

    monkeypatch.setattr("docaudit.ml.components.new_source_id", mock_new_source_id)
    assert list(iter_sources(sources, source_ids)) == expected


@pytest.mark.parametrize(
    "d1,d2,expected",
    [
        # General assumption: Values of d1 have priority
        # Case 1: Empty dictionaries
        ({}, {}, {}),
        # Case 2: No overlapping keys
        ({"a": 1}, {"b": 2}, {"a": 1, "b": 2}),
        # Case 3: Overlapping keys, no deeper structures
        ({"a": 1}, {"a": 2}, {"a": 1}),
        # Case 4: Nested dictionaries
        ({"a": {"b": 1}}, {"a": {"b": 2}}, {"a": {"b": 1}}),
        # Case 5: Lists as values
        ({"a": [1, 2]}, {"a": [3, 4]}, {"a": [1, 2, 3, 4]}),
        # Case 6: Lists with dictionaries as values
        ({"a": [{"b": 1}]}, {"a": [{"b": 2}]}, {"a": [{"b": 1}, {"b": 2}]}),
        # Case 7: Unequal structures for equal keys, reversion to d1 value
        ({"a": [1, 2]}, {"a": {"b": 2}}, {"a": [1, 2]}),
        # Case 8: A key only in d2
        ({}, {"a": 2}, {"a": 2}),
    ],
)
def test_recursively_merge_dicts(d1, d2, expected):
    assert recursively_merge_dicts(d1, d2) == expected
