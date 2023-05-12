# Copyright (C) 2023 Helmar Hutschenreuter
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <http://www.gnu.org/licenses/>.

import re
from typing import Iterator

import docx
from docx.text.paragraph import Paragraph
from haystack.schema import Document


def parse_level(paragraph: Paragraph) -> int | None:
    style = paragraph.style.name
    match = re.search(r"Heading (\d+)", style)
    if match:
        return int(match.group(1))
    return None


def modify_headers(paragraph, headers: list[str]) -> bool:
    level = parse_level(paragraph)
    if level is None:
        return False

    if len(headers) < level:
        headers.append(paragraph.text)
        return True

    while len(headers) > level:
        headers.pop()

    # len(headers) == level
    headers.pop()
    headers.append(paragraph.text)
    return True


def to_haystack_document(headers, content) -> Document:
    return Document(
        meta=dict(headers=headers),
        content=content[:-2],
        id_hash_keys=["meta", "content"],
    )


def parse_docx(filename: str) -> Iterator[Document]:
    headers = []
    content = ""
    for paragraph in docx.Document(filename).paragraphs:
        if modify_headers(paragraph, headers):
            yield to_haystack_document(headers, content)

            # Start with a new section
            content = paragraph.text + "\n" * 2
        else:
            # Continue with current section
            content += paragraph.text + "\n" * 2

    yield to_haystack_document(headers, content)
