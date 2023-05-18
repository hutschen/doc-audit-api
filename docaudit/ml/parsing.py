# Copyright (C) 2023 Helmar Hutschenreuter
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <http://www.gnu.org/licenses/>.

import re
from typing import Iterator, List

import docx
from docx.text.paragraph import Paragraph
from haystack.nodes import BaseComponent
from haystack.schema import Document


class DocxParser(BaseComponent):
    outgoing_edges = 1

    @staticmethod
    def _parse_level(paragraph: Paragraph) -> int | None:
        style = paragraph.style.name
        match = re.search(r"Heading (\d+)", style)
        if match:
            return int(match.group(1))
        return None

    @classmethod
    def _modify_headers(cls, paragraph, headers: list[str]) -> bool:
        level = cls._parse_level(paragraph)
        if level is None:
            return False

        if len(headers) < level:
            headers.append(paragraph.text)
            return True

        while len(headers) > level:
            headers.pop()

        # len(headers) == level
        headers.pop()
        headers.append(paragraph.text)
        return True

    @staticmethod
    def _to_document(headers, content, meta: dict | None = None) -> Document:
        return Document(
            # Combine meta and headers into one meta dict
            meta={
                **(meta or {}),
                "headers": headers,
            },
            content=content[:-2],
            id_hash_keys=["meta", "content"],
        )

    @classmethod
    def _parse_docx(cls, filename: str, meta: dict | None = None) -> Iterator[Document]:
        headers = []
        content = ""
        for paragraph in docx.Document(filename).paragraphs:
            prev_headers = headers.copy()
            if cls._modify_headers(paragraph, headers):
                # New section started, yield the previous section
                yield cls._to_document(prev_headers, content, meta)

                # Start with a new section
                content = paragraph.text + "\n" * 2
            else:
                # Continue with current section
                content += paragraph.text + "\n" * 2

        yield cls._to_document(headers, content, meta)

    def __init__(self):
        pass

    def run(self, *, file_paths: List[str], meta: dict | None = None, **kwargs):
        documents = []
        for file_path in file_paths:
            documents.extend(self._parse_docx(file_path, meta))
        output = {"documents": documents, **kwargs}
        return output, "output_1"

    def run_batch(self, **_):
        raise NotImplementedError("run_batch is not implemented for DocxParser")
