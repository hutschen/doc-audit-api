# Copyright (C) 2024 Helmar Hutschenreuter
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

import hashlib
import re
import uuid
from typing import IO, Any, Dict, Generator, List

import docx
from docx.text.paragraph import Paragraph
from haystack import Document, component, logging
from haystack.core.component.types import Variadic
from haystack.document_stores.types import DocumentStore

from ..utils import remove_extra_whitespace

logger = logging.getLogger(__name__)


class DocxParserError(Exception):
    pass


class DocxParser:
    @staticmethod
    def parse_level(paragraph: Paragraph) -> int | None:
        """
        Parses the level of a heading paragraph.
        If the paragraph is not a heading, returns None.
        """
        style = paragraph.style.name
        match = re.search(r"Heading (\d+)", style)
        return int(match.group(1)) if match else None

    @staticmethod
    def open_docx_carefully(docx_input: str | IO[bytes]) -> docx.Document:
        """
        Carefully opens a DOCX file and raises a DocxParserError if it fails.
        """
        try:
            return docx.Document(docx_input)
        except Exception as e:
            raise DocxParserError(f"Failed to open DOCX file. Error: {e}") from e

    @classmethod
    def parse(
        cls, docx_input: str | IO[bytes]
    ) -> Generator[tuple[tuple, str] | tuple[tuple, str], Any, None]:
        """
        Parses the headers and content from a DOCX file.
        """
        headers = []
        contents = []

        for paragraph in cls.open_docx_carefully(docx_input).paragraphs:
            paragraph_level = cls.parse_level(paragraph)
            paragraph_text = remove_extra_whitespace(paragraph.text)
            if paragraph_level is not None:
                # Paragraph is a heading, so yield the previous content
                yield tuple(headers), ("\n" * 2).join(contents)

                # Update headers and reset content
                headers = headers[: paragraph_level - 1]
                headers.append(paragraph_text)
                contents = [paragraph_text]
            else:
                # Paragraph is not a heading
                contents.append(paragraph_text)

        yield tuple(headers), ("\n" * 2).join(contents)


def new_source_id() -> str:
    return str(uuid.uuid4())


def iter_sources(sources: List[str | IO[bytes]], source_ids: List[str] | None):
    source_ids = [] if source_ids is None else source_ids
    length_diff = len(sources) - len(source_ids)
    if length_diff > 0:
        source_ids += [new_source_id() for _ in range(length_diff)]
    elif length_diff < 0:
        source_ids = source_ids[: len(sources)]
    return ((s, sid) for s, sid in zip(sources, source_ids))


@component
class DocxToDocuments:
    """
    Converts DOCX files to Documents.
    """

    @component.output_types(documents=List[Document])
    def run(
        self,
        sources: List[str | IO[bytes]],
        source_ids: List[str] | None = None,
        meta: Dict[str, Any] | None = None,
    ) -> Dict[str, List[Document]]:
        """
        Converts a list of DOCX files to Documents.
        """

        def generate_documents():
            for source, source_id in iter_sources(sources, source_ids):
                try:
                    for headers, content in DocxParser.parse(source):
                        yield Document(
                            meta={
                                **(meta or {}),
                                "locations": [
                                    {
                                        "id": source_id,
                                        "type": "docx",
                                        "path": headers,
                                    }
                                ],
                            },
                            content=content,
                        )
                except DocxParserError as e:
                    logger.warning(
                        "Failed to convert DOCX file. Skipping it. Error: {error}",
                        error=e,
                    )

        return {"documents": list(generate_documents())}


@component
class SetContentBasedIds:
    """
    Sets the IDs of documents based on their content. WARNING: This can result in
    duplicate IDs when the content is the same but the metadata is different.
    """

    @component.output_types(documents=List[Document])
    def run(
        self,
        documents: List[Document],
    ) -> Dict[str, List[Document]]:
        """
        Sets the IDs of documents based on their content.
        """

        for doc in documents:
            doc.id = hashlib.sha256(doc.content.encode("utf-8")).hexdigest()

        return {"documents": documents}


def recursively_merge_dicts(d1: Dict, d2: Dict) -> Dict:
    """
    Recursively merges two dictionaries.
    """
    merged = {}
    all_keys = d1.keys() | d2.keys()

    for key in all_keys:
        if key in d1 and key in d2:
            if isinstance(d1[key], dict) and isinstance(d2[key], dict):
                merged[key] = recursively_merge_dicts(d1[key], d2[key])
            elif isinstance(d1[key], list) and isinstance(d2[key], list):
                merged[key] = d1[key] + d2[key]
            else:
                merged[key] = d1[key]
        elif key in d1:
            merged[key] = d1[key]
        else:  # Key is only in d2
            merged[key] = d2[key]

    return merged


@component
class MergeMetadata:
    """
    Merges the metadata of documents with the same ID. WARNING: This results in removing
    of duplicate documents.
    """

    @component.output_types(documents=List[Document])
    def run(self, documents: Variadic[List[Document]]) -> Dict[str, List[Document]]:
        """
        Merges the metadata of documents with the same ID.
        """

        # Create a dictionary mapping IDs to documents
        id_to_documents: Dict[str, List[Document]] = {}
        for documents_list in documents:
            for doc in documents_list:
                id_to_documents.setdefault(doc.id, []).append(doc)

        # Merge the metadata of documents with the same ID
        # Keep the first document and merge the metadata of the rest into it
        merged_documents = []
        for docs in id_to_documents.values():
            merged_doc = docs[0]
            for doc in docs[1:]:
                merged_doc.meta = recursively_merge_dicts(merged_doc.meta, doc.meta)
            merged_documents.append(merged_doc)

        return {"documents": merged_documents}


@component
class DuplicateChecker:
    """
    Checks if documents with the same ID already exist in the Document Store. To
    optimize performance of remote  stores, queries to the store are batched.
    """

    def __init__(self, document_store: DocumentStore, batch_size: int = 32):
        """
        Create an DuplicatesChecker component.

        :param document_store:
            Document store to check.
        :param batch_size:
            Number of documents to check in at once against the Document Store.
        """
        self.document_store = document_store
        self.batch_size = batch_size

    @component.output_types(
        retrieved=List[Document], hits=List[Document], misses=List[Document]
    )
    def run(self, documents: List[Document]) -> Dict[str, List[Document]]:
        """
        Checks if documents with the same ID already exist in the Document Store.
        """
        retrieved = []
        hits = []
        misses = []

        for i in range(0, len(documents), self.batch_size):
            batch = documents[i : i + self.batch_size]
            existing_docs = self.document_store.filter_documents(
                dict(field="id", operator="in", value=[doc.id for doc in batch])
            )
            retrieved.extend(existing_docs)

            for doc in batch:
                if any(existing_doc.id == doc.id for existing_doc in existing_docs):
                    hits.append(doc)
                else:
                    misses.append(doc)

        return {"retrieved": retrieved, "hits": hits, "misses": misses}


@component
class LocationRemover:
    """
    Remove locations from the metadata of documents.
    """

    @component.output_types(documents=List[Document])
    def run(
        self, documents: List[Document], source_ids: List[str] | None = None
    ) -> Dict[str, List[Document]]:
        """
        Remove all locations associated with specific source IDs from the metadata of
        documents.
        """

        if source_ids is None:
            return {"documents": documents}

        for doc in documents:
            doc.meta["locations"] = [
                location
                for location in doc.meta.get("locations", [])
                if location["id"] not in source_ids
            ]

        return {"documents": documents}
